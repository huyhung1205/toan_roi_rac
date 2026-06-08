from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_FILE = Path("Ly_thuyet_Thi_cuoi_ky_Toan_roi_rac.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_paragraph_spacing(p, before=0, after=0, line=1.15):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_run(p, text, bold=False, italic=False, color=None, size=None):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor(*color)
    if size:
        r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    return r


def add_heading(document, text, level=1):
    p = document.add_paragraph()
    p.style = document.styles[f"Heading {level}"]
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    if level == 1:
        r.font.size = Pt(15)
        r.font.color.rgb = RGBColor(31, 78, 121)
    else:
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(46, 125, 50)
    set_paragraph_spacing(p, before=8, after=4, line=1.1)
    return p


def add_body(document, text, bold_prefix=None):
    p = document.add_paragraph()
    set_paragraph_spacing(p, after=2)
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p, bold_prefix, bold=True, color=(55, 71, 79))
        add_run(p, text[len(bold_prefix):], color=(38, 50, 56))
    else:
        add_run(p, text, color=(38, 50, 56))
    return p


def add_bullet(document, text, level=0):
    p = document.add_paragraph(style="List Bullet")
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.25 * level)
    set_paragraph_spacing(p, after=1)
    add_run(p, text, color=(38, 50, 56))
    return p


def add_number(document, text, level=0):
    p = document.add_paragraph(style="List Number")
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.25 * level)
    set_paragraph_spacing(p, after=1)
    add_run(p, text, color=(38, 50, 56))
    return p


def add_note_box(document, title, lines, fill="EAF4F4", title_color=(20, 80, 86)):
    tbl = document.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=2)
    add_run(p, title, bold=True, color=title_color, size=11.5)
    for line in lines:
        pp = cell.add_paragraph()
        pp.style = document.styles["List Bullet"]
        set_paragraph_spacing(pp, after=0)
        add_run(pp, line, color=(38, 50, 56), size=10.5)
    document.add_paragraph()


def add_table(document, headers, rows, col_widths=None, header_fill="DCEAF5"):
    tbl = document.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        set_cell_margins(hdr[i])
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(31, 78, 121)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            set_cell_margins(cells[i])
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(10)
                    r.font.color.rgb = RGBColor(38, 50, 56)
    if col_widths:
        for row in tbl.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)
    document.add_paragraph()
    return tbl


def set_doc_defaults(document):
    sec = document.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Times New Roman"


def add_title_page(document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=4)
    add_run(p, "TÀI LIỆU ÔN THI CUỐI KỲ", bold=True, color=(31, 78, 121), size=22)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=2)
    add_run(p, "TOÁN RỜI RẠC", bold=True, color=(46, 125, 50), size=24)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, after=10)
    add_run(p, "Tổng hợp lý thuyết từ B5 đến B12, bám sát các dạng bài trong main.pdf", italic=True, color=(90, 90, 90), size=11)

    add_note_box(
        document,
        "Mục tiêu của tài liệu",
        [
            "Nhận dạng nhanh đúng dạng bài trong đề.",
            "Biết công thức, định lý, và quy trình làm bài chuẩn.",
            "Có thể xử lý các bài từ main.pdf theo từng bước, không học lan man.",
        ],
        fill="EEF6EE",
        title_color=(46, 125, 50),
    )

    add_note_box(
        document,
        "Quy ước khi làm bài trong đề",
        [
            "Nếu ma trận kề đối xứng thì đồ thị vô hướng; nếu bất đối xứng thì đồ thị có hướng.",
            "Nếu ma trận không có trọng số thì có thể xem mỗi cạnh/cung có trọng số bằng 1.",
            "Khi DFS, BFS, Kruskal, Prim có nhiều lựa chọn cùng tốt thì lấy theo thứ tự chữ cái.",
        ],
        fill="F4F7FB",
        title_color=(31, 78, 121),
    )


def build_document():
    doc = Document()
    set_doc_defaults(doc)
    add_title_page(doc)

    add_heading(doc, "1. Bản đồ đề thi nên học trước", 1)
    add_body(doc, "Các bài trong main.pdf xoay quanh 7 nhóm kỹ năng sau:")
    for item in [
        "Đọc ma trận kề / ma trận trọng số để dựng lại đồ thị.",
        "Duyệt DFS, BFS từ một đỉnh cho trước.",
        "Kiểm tra Euler, chu trình Euler, đường đi Euler.",
        "Tìm cây khung nhỏ nhất bằng Kruskal và Prim.",
        "Tìm đường đi ngắn nhất bằng Dijkstra, Bellman-Ford, Floyd.",
        "Kiểm tra và xây dựng luồng cực đại bằng Ford-Fulkerson.",
        "Lập hàm Boole từ bảng chân trị và tối thiểu hóa bằng Karnaugh.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2. B5 - Các khái niệm cơ bản về đồ thị", 1)
    add_body(doc, "2.1. Định nghĩa:")
    add_bullet(doc, "Đồ thị là một cấu trúc rời rạc gồm tập đỉnh V và tập cạnh E nối các đỉnh.")
    add_bullet(doc, "Đồ thị vô hướng: cạnh là một bộ không có thứ tự (u, v).")
    add_bullet(doc, "Đồ thị có hướng: cung là một bộ có thứ tự (u, v), trong đó u là đỉnh đầu, v là đỉnh cuối.")
    add_bullet(doc, "Đa đồ thị cho phép cạnh song song; giả đồ thị cho phép khuyên (cạnh nối một đỉnh với chính nó).")

    add_body(doc, "2.2. Thuật ngữ cơ bản:")
    add_bullet(doc, "Hai đỉnh kề nhau nếu có cạnh nối chúng.")
    add_bullet(doc, "Một cạnh liên thuộc với hai đầu mút của nó.")
    add_bullet(doc, "Đỉnh cô lập: bậc 0. Đỉnh treo: bậc 1.")
    add_bullet(doc, "Đồ thị con: lấy một tập đỉnh và một tập cạnh con từ đồ thị gốc.")
    add_bullet(doc, "Đường đi: dãy đỉnh liên tiếp nối bởi các cạnh. Chu trình: đường đi khép kín.")
    add_bullet(doc, "Đồ thị liên thông: giữa mọi cặp đỉnh đều có đường đi.")

    add_body(doc, "2.3. Bậc đỉnh và định lý bắt tay:")
    add_bullet(doc, "Đồ thị vô hướng: deg(v) là số cạnh kề với v.")
    add_bullet(doc, "Đồ thị có hướng: deg+(v) là bán bậc ra, deg-(v) là bán bậc vào.")
    add_bullet(doc, "Định lý bắt tay: tổng bậc của mọi đỉnh trong đồ thị vô hướng bằng 2|E|.")
    add_bullet(doc, "Hệ quả: số đỉnh bậc lẻ trong một đồ thị vô hướng luôn là số chẵn.")
    add_bullet(doc, "Với đồ thị có hướng: tổng tất cả bán bậc ra bằng tổng tất cả bán bậc vào bằng số cung.")

    add_note_box(
        doc,
        "Mẹo đọc đề",
        [
            "Nếu đề yêu cầu kiểm tra khả thi bằng định lý bắt tay, chỉ cần tính tổng bậc và xét số đỉnh bậc lẻ.",
            "Nếu đồ thị có hướng, luôn tách riêng deg+ và deg-.",
        ],
        fill="FFF7E8",
        title_color=(140, 92, 0),
    )

    add_heading(doc, "3. B6 - Biểu diễn đồ thị", 1)
    add_body(doc, "Các cách biểu diễn chính:")
    add_bullet(doc, "Ma trận kề A = [aij]. Với đồ thị vô hướng: aij = 1 nếu có cạnh, 0 nếu không có.")
    add_bullet(doc, "Ma trận trọng số C = [cij]. Nếu không có cạnh thì thường dùng 0, +∞ hoặc một ký hiệu quy ước θ.")
    add_bullet(doc, "Ma trận liên thuộc đỉnh-cạnh: ghi quan hệ giữa đỉnh và cạnh; với đồ thị có hướng thường dùng +1 ở đỉnh đầu, -1 ở đỉnh cuối.")
    add_bullet(doc, "Danh sách kề Ke[v]: lưu các đỉnh kề với v.")
    add_bullet(doc, "Danh sách cạnh: lưu toàn bộ cạnh/cung theo thứ tự.")

    add_table(
        doc,
        ["Biểu diễn", "Ưu điểm", "Nhược điểm", "Khi nào dùng"],
        [
            ["Ma trận kề", "Kiểm tra hai đỉnh kề nhau rất nhanh.", "Tốn O(n²) bộ nhớ, duyệt hàng xóm chậm.", "Đồ thị dày, cần tra cứu cạnh nhanh."],
            ["Danh sách kề", "Tiết kiệm bộ nhớ O(n+m), tốt cho duyệt BFS/DFS.", "Kiểm tra kề nhau không nhanh bằng ma trận.", "Đồ thị thưa, thuật toán duyệt."],
            ["Danh sách cạnh", "Rất gọn, thuận lợi cho Kruskal.", "Tìm hàng xóm phải quét toàn bộ cạnh.", "Bài toán cây khung, đồ thị thưa."],
        ],
        col_widths=[1.4, 2.2, 2.2, 1.5],
    )

    add_body(doc, "Tính chất cần nhớ của ma trận kề:")
    add_bullet(doc, "Đồ thị vô hướng: ma trận đối xứng A = AT.")
    add_bullet(doc, "Tổng mọi phần tử của A bằng 2m với m là số cạnh.")
    add_bullet(doc, "deg(vk) = tổng dòng k = tổng cột k.")
    add_bullet(doc, "Đồ thị có hướng: tổng mọi phần tử bằng m.")
    add_bullet(doc, "deg+(vk) = tổng dòng k; deg-(vk) = tổng cột k.")
    add_bullet(doc, "Nếu A^p = [āij] thì āij là số đường đi khác nhau từ vi đến vj qua p-1 đỉnh trung gian.")

    add_body(doc, "Cách làm bài từ ma trận ra đồ thị:")
    add_bullet(doc, "Đọc thứ tự hàng và cột để xác định tập đỉnh.")
    add_bullet(doc, "Ma trận đối xứng thường là đồ thị vô hướng; ma trận không đối xứng thường là đồ thị có hướng.")
    add_bullet(doc, "Nếu chỉ có 0 và 1 thì là ma trận kề không trọng số; nếu có số khác 0, 1 thì đó là trọng số.")
    add_bullet(doc, "Duyệt từng ô aij: nếu aij = 0 thì không vẽ; nếu aij != 0 thì vẽ cạnh/cung tương ứng.")
    add_bullet(doc, "Với đồ thị vô hướng, chỉ cần xét nửa trên hoặc nửa dưới để tránh vẽ lặp.")
    add_bullet(doc, "Với đồ thị có hướng, mỗi ô aij tương ứng đúng một cung vi -> vj.")
    add_bullet(doc, "Ghi trọng số lên cạnh/cung nếu đề cho trọng số.")
    add_bullet(doc, "Đếm lại số cạnh/cung nếu đề yêu cầu kiểm tra kết quả.")

    add_body(doc, "Cách làm bài từ đồ thị ra ma trận:")
    add_bullet(doc, "Chọn thứ tự đỉnh cố định, thường là A, B, C, ...")
    add_bullet(doc, "Lập ma trận theo đúng thứ tự đó.")
    add_bullet(doc, "Với đồ thị vô hướng, nếu có cạnh giữa vi và vj thì ghi vào cả aij và aji.")
    add_bullet(doc, "Với đồ thị có hướng, nếu có cung vi -> vj thì ghi vào aij.")
    add_bullet(doc, "Nếu không có cạnh/cung thì ghi 0; đường chéo chính thường là 0 nếu không có khuyên.")

    add_heading(doc, "4. B7 - Thuật toán duyệt đồ thị: BFS và DFS", 1)
    add_body(doc, "4.1. BFS - duyệt theo chiều rộng:")
    add_bullet(doc, "Cấu trúc dữ liệu: hàng đợi Queue, nguyên tắc FIFO.")
    add_bullet(doc, "BFS duyệt theo từng mức: đỉnh xuất hiện sớm sẽ được xử lý sớm.")
    add_bullet(doc, "Trong đồ thị không trọng số, mức của một đỉnh chính là độ dài đường đi ngắn nhất theo số cạnh từ đỉnh xuất phát.")
    add_bullet(doc, "Phù hợp để tìm đường đi ngắn nhất theo số cạnh, kiểm tra liên thông, tạo cây BFS.")

    add_body(doc, "4.2. DFS - duyệt theo chiều sâu:")
    add_bullet(doc, "Cấu trúc dữ liệu: stack hoặc đệ quy.")
    add_bullet(doc, "DFS đi sâu tối đa có thể trước khi quay lui.")
    add_bullet(doc, "Phù hợp để kiểm tra thành phần liên thông, phát hiện chu trình, sinh cây DFS.")

    add_table(
        doc,
        ["Thuật toán", "Cấu trúc", "Ý nghĩa", "Độ phức tạp"],
        [
            ["BFS", "Queue", "Duyệt theo mức; tìm đường đi ngắn nhất theo số cạnh.", "O(n + m) với danh sách kề"],
            ["DFS", "Stack / đệ quy", "Duyệt đi sâu; hữu ích cho kiểm tra cấu trúc đồ thị.", "O(n + m) với danh sách kề"],
        ],
        col_widths=[1.0, 1.0, 3.4, 1.5],
    )

    add_note_box(
        doc,
        "Mẫu trình bày khi làm bài BFS/DFS",
        [
            "Bước 1: Xác định danh sách kề theo thứ tự chữ cái.",
            "Bước 2: Ghi cây duyệt hoặc thứ tự thăm đỉnh.",
            "Bước 3: Không bỏ qua đỉnh đã thăm; chỉ thêm đỉnh chưa thăm vào hàng đợi/ngăn xếp.",
        ],
        fill="F4F7FB",
        title_color=(31, 78, 121),
    )

    add_heading(doc, "5. B8 - Đồ thị Euler và Hamilton", 1)
    add_body(doc, "5.1. Euler:")
    add_bullet(doc, "Đường đi Euler: đi qua tất cả các cạnh của đồ thị, mỗi cạnh đúng một lần.")
    add_bullet(doc, "Chu trình Euler: là đường đi Euler khép kín.")
    add_bullet(doc, "Đồ thị nửa Euler: có đường đi Euler nhưng không có chu trình Euler.")
    add_bullet(doc, "Đồ thị Euler: có chu trình Euler.")
    add_bullet(doc, "Với đồ thị có hướng, điều kiện được xét trên bán bậc vào và bán bậc ra.")

    add_table(
        doc,
        ["Loại đồ thị", "Điều kiện có chu trình Euler", "Điều kiện có đường đi Euler"],
        [
            ["Vô hướng liên thông", "Mọi đỉnh đều có bậc chẵn.", "Có đúng 2 đỉnh bậc lẻ."],
            ["Có hướng liên thông yếu", "Mọi đỉnh có deg-(v) = deg+(v).", "Tồn tại đúng 1 đỉnh deg+ - deg- = 1 và đúng 1 đỉnh deg- - deg+ = 1; các đỉnh còn lại cân bằng."],
        ],
        col_widths=[1.6, 2.8, 2.8],
    )

    add_body(doc, "5.2. Quy trình kiểm tra Euler:")
    add_bullet(doc, "Bước 1: Kiểm tra liên thông hoặc liên thông yếu.")
    add_bullet(doc, "Bước 2: Tính bậc / bán bậc của từng đỉnh.")
    add_bullet(doc, "Bước 3: So sánh với điều kiện Euler hoặc nửa Euler.")
    add_bullet(doc, "Bước 4: Nếu cần dựng đường đi, dùng stack theo kiểu Hierholzer.")

    add_body(doc, "5.3. Hamilton:")
    add_bullet(doc, "Đường đi Hamilton: đi qua mỗi đỉnh đúng một lần.")
    add_bullet(doc, "Chu trình Hamilton: là đường đi Hamilton khép kín.")
    add_bullet(doc, "Khác với Euler, Hamilton không có một tiêu chuẩn đơn giản, phổ quát như bậc chẵn.")
    add_bullet(doc, "Khi làm bài cơ bản, thường phải kiểm tra trực tiếp bằng cách xây dựng chu trình/đường đi.")
    add_bullet(doc, "Nếu môn học cho phép dùng định lý đủ điều kiện, có thể nhớ thêm: Dirac và Ore là các điều kiện đủ quen thuộc cho Hamilton.")

    add_heading(doc, "6. B9 - Cây và cây khung", 1)
    add_body(doc, "6.1. Cây:")
    add_bullet(doc, "Cây là đồ thị vô hướng, liên thông và không có chu trình.")
    add_bullet(doc, "Rừng là đồ thị mà mỗi thành phần liên thông là một cây.")
    add_bullet(doc, "Một cây n đỉnh có đúng n - 1 cạnh.")
    add_bullet(doc, "Giữa hai đỉnh bất kỳ của cây luôn có đúng một đường đi đơn.")

    add_body(doc, "6.2. Cây khung:")
    add_bullet(doc, "Cây khung của đồ thị liên thông G là một cây chứa tất cả các đỉnh của G.")
    add_bullet(doc, "Nếu đồ thị có trọng số, cây khung nhỏ nhất là cây khung có tổng trọng số nhỏ nhất.")

    add_body(doc, "6.3. Kruskal:")
    add_bullet(doc, "Sắp xếp các cạnh theo trọng số tăng dần.")
    add_bullet(doc, "Lần lượt chọn cạnh nhỏ nhất chưa tạo chu trình.")
    add_bullet(doc, "Dừng khi có n - 1 cạnh.")
    add_bullet(doc, "Nếu có cạnh cùng trọng số, dùng quy ước chữ cái để phá hòa.")

    add_body(doc, "6.4. Prim:")
    add_bullet(doc, "Bắt đầu từ một đỉnh bất kỳ.")
    add_bullet(doc, "Mỗi bước chọn cạnh nhỏ nhất nối tập đỉnh đã chọn với một đỉnh chưa chọn.")
    add_bullet(doc, "Tiếp tục đến khi có n - 1 cạnh.")

    add_table(
        doc,
        ["Thuật toán", "Ý tưởng chính", "Phù hợp với", "Độ phức tạp thường dùng"],
        [
            ["Kruskal", "Chọn cạnh nhỏ nhất có thể mà không tạo chu trình.", "Danh sách cạnh; đồ thị thưa.", "O(m log m) do sắp xếp"],
            ["Prim", "Mở rộng từ một đỉnh theo cạnh nhỏ nhất ra ngoài cây.", "Ma trận kề / danh sách kề.", "O(m log n) hoặc O(n²) tùy cài đặt"],
        ],
        col_widths=[1.1, 2.9, 1.5, 1.6],
    )

    add_heading(doc, "7. B10 - Đường đi ngắn nhất", 1)
    add_body(doc, "7.1. Khái niệm:")
    add_bullet(doc, "Trọng số của một đường đi là tổng trọng số các cạnh trên đường đi đó.")
    add_bullet(doc, "Đường đi ngắn nhất từ u đến v là đường đi có tổng trọng số nhỏ nhất.")
    add_bullet(doc, "Ký hiệu khoảng cách ngắn nhất: δ(u, v).")
    add_bullet(doc, "Đường đi ngắn nhất luôn có thể chọn là đường đi đơn.")
    add_bullet(doc, "Đường đi ngắn nhất không quá n - 1 cạnh.")
    add_bullet(doc, "Mọi đoạn con của đường đi ngắn nhất cũng là đường đi ngắn nhất.")

    add_body(doc, "7.2. Relaxation - giảm cận trên:")
    add_bullet(doc, "Lưu d[v]: độ dài tốt nhất hiện biết từ s đến v.")
    add_bullet(doc, "Lưu p[v]: đỉnh trước v trên đường đi tốt nhất.")
    add_bullet(doc, "Relax(u, v): nếu d[v] > d[u] + w(u, v) thì cập nhật d[v] và p[v].")

    add_body(doc, "7.3. Bellman-Ford:")
    add_bullet(doc, "Dùng được cho trọng số âm, miễn là không có chu trình âm.")
    add_bullet(doc, "Khởi tạo d[s] = 0; các đỉnh khác là vô cùng.")
    add_bullet(doc, "Lặp relax toàn bộ cạnh |V| - 1 lần.")
    add_bullet(doc, "Sau lượt lặp cuối, nếu còn relax được thì đồ thị có chu trình âm.")
    add_bullet(doc, "Hữu ích khi cần xử lý bài có cạnh âm hoặc kiểm tra chu trình âm.")

    add_body(doc, "7.4. Dijkstra:")
    add_bullet(doc, "Chỉ áp dụng khi mọi trọng số không âm.")
    add_bullet(doc, "Mỗi bước chọn đỉnh chưa cố định có d nhỏ nhất, sau đó relax các cạnh đi ra từ đỉnh đó.")
    add_bullet(doc, "Khi một đỉnh được chọn, nhãn của nó trở thành cố định.")
    add_bullet(doc, "Rất phù hợp cho các bài trong main.pdf vì đa số trọng số đều không âm.")

    add_body(doc, "7.5. Floyd-Warshall:")
    add_bullet(doc, "Giải bài toán mọi cặp đỉnh.")
    add_bullet(doc, "Cập nhật theo công thức: dist[i][j] = min(dist[i][j], dist[i][k] + dist[k][j]).")
    add_bullet(doc, "Dùng thêm ma trận đường đi để truy vết lại toàn bộ đường đi ngắn nhất.")

    add_table(
        doc,
        ["Thuật toán", "Bài toán", "Điều kiện", "Ghi nhớ nhanh"],
        [
            ["BFS", "Đường đi ngắn nhất theo số cạnh", "Đồ thị không trọng số", "Theo từng mức"],
            ["Dijkstra", "Đường đi ngắn nhất một nguồn", "Không có trọng số âm", "Chọn đỉnh có d nhỏ nhất"],
            ["Bellman-Ford", "Đường đi ngắn nhất một nguồn", "Cho phép trọng số âm, không có chu trình âm", "Relax toàn bộ cạnh nhiều vòng"],
            ["Floyd", "Mọi cặp đỉnh", "Không có chu trình âm nếu cần truy vết ổn định", "3 vòng lặp i-j-k"],
        ],
        col_widths=[1.15, 1.9, 2.1, 1.8],
    )

    add_heading(doc, "8. B11 - Luồng cực đại", 1)
    add_body(doc, "8.1. Mạng và luồng:")
    add_bullet(doc, "Mạng là đồ thị có hướng có duy nhất một nguồn s và một đích t.")
    add_bullet(doc, "Mỗi cung e có khả năng thông qua c(e) >= 0.")
    add_bullet(doc, "Luồng f thỏa 0 <= f(e) <= c(e).")
    add_bullet(doc, "Tại mọi đỉnh trung gian: tổng luồng vào = tổng luồng ra.")
    add_bullet(doc, "Giá trị luồng: val(f) = tổng luồng ra khỏi nguồn = tổng luồng vào đích.")

    add_body(doc, "8.2. Lát cắt:")
    add_bullet(doc, "Lát cắt (S, T) là phân hoạch đỉnh sao cho s ∈ S và t ∈ T.")
    add_bullet(doc, "Khả năng thông qua của lát cắt: cap(S, T) = tổng c(e) của các cung đi từ S sang T.")
    add_bullet(doc, "Luồng qua lát cắt bằng tổng luồng đi từ S sang T trừ tổng luồng ngược từ T sang S.")
    add_bullet(doc, "Định lý: val(f) <= cap(S, T) với mọi lát cắt.")
    add_bullet(doc, "Nếu val(f) = cap(S, T) thì luồng là cực đại và lát cắt là lát cắt nhỏ nhất.")

    add_body(doc, "8.3. Đồ thị tăng luồng và đường tăng luồng:")
    add_bullet(doc, "Đồ thị tăng luồng cho biết phần dung lượng còn có thể đẩy thêm.")
    add_bullet(doc, "Cung thuận có dung lượng c - f; cung nghịch có dung lượng f.")
    add_bullet(doc, "Đường tăng luồng là đường từ s đến t trong đồ thị tăng luồng.")
    add_bullet(doc, "Dung lượng của đường tăng luồng là giá trị nhỏ nhất trên các cung của đường.")

    add_body(doc, "8.4. Ford-Fulkerson:")
    add_bullet(doc, "Khởi tạo luồng bằng 0.")
    add_bullet(doc, "Liên tục tìm đường tăng luồng và tăng luồng theo dung lượng nút cổ chai.")
    add_bullet(doc, "Dừng khi không còn đường tăng luồng.")
    add_bullet(doc, "Theo định lý Ford-Fulkerson: luồng cực đại khi và chỉ khi không còn đường tăng luồng.")

    add_note_box(
        doc,
        "Quy trình làm bài max-flow",
        [
            "Vẽ mạng và ghi dung lượng trên từng cung.",
            "Lập đồ thị dư.",
            "Tìm một đường tăng luồng từ s đến t.",
            "Lấy nút cổ chai, cập nhật luồng và đồ thị dư.",
            "Lặp lại đến khi bế tắc.",
        ],
        fill="F4F7FB",
        title_color=(31, 78, 121),
    )

    add_heading(doc, "9. B12 - Đại số Boole", 1)
    add_body(doc, "9.1. Đại số Boole:")
    add_bullet(doc, "Đại số Boole là hệ (B, ∧, ∨, ¬, 0, 1) thỏa mãn các tiên đề giao hoán, kết hợp, phân phối, đơn vị và bù.")
    add_bullet(doc, "Trong thực hành: x + y tương ứng với x ∨ y; xy tương ứng với x ∧ y; x̄ tương ứng với ¬x.")
    add_bullet(doc, "Thứ tự tính: ¬ trước, rồi ∧, cuối cùng ∨.")

    add_body(doc, "9.2. Các tính chất cơ bản cần nhớ:")
    add_bullet(doc, "Nuốt: a ∧ 0 = 0, a ∨ 1 = 1.")
    add_bullet(doc, "Lũy đẳng: a ∨ a = a, a ∧ a = a.")
    add_bullet(doc, "Hấp thụ: a ∨ (a ∧ b) = a, a ∧ (a ∨ b) = a.")
    add_bullet(doc, "Bù kép: ¬(¬a) = a.")
    add_bullet(doc, "De Morgan: ¬(a ∨ b) = ¬a ∧ ¬b; ¬(a ∧ b) = ¬a ∨ ¬b.")

    add_body(doc, "9.3. Hàm Boole và dạng chuẩn tắc:")
    add_bullet(doc, "Hàm Boole n biến là ánh xạ f: B^n -> B.")
    add_bullet(doc, "Từ bảng chân trị, các dòng có f = 1 sinh ra các minterm, cộng tất cả lại để được dạng tổng chuẩn tắc đầy đủ.")
    add_bullet(doc, "Với một dòng có x = 0 thì lấy x̄, có x = 1 thì lấy x; sau đó nhân tất cả các biến của dòng đó.")
    add_bullet(doc, "Dạng tích chuẩn tắc có thể xây dựng tương tự từ các dòng có f = 0.")

    add_body(doc, "9.4. Rút gọn và tối thiểu hóa:")
    add_bullet(doc, "Dùng các luật đại số để rút gọn trực tiếp khi biểu thức ngắn.")
    add_bullet(doc, "Dùng bản đồ Karnaugh khi cần tối thiểu hóa hệ thống nhiều biến.")
    add_bullet(doc, "Nhóm các ô 1 thành hình chữ nhật kích thước 1, 2, 4, 8,...; nhóm càng lớn càng tốt.")
    add_bullet(doc, "Ô ở mép trái và mép phải, hoặc mép trên và mép dưới, được phép coi là kề nhau theo quy tắc quấn vòng.")
    add_bullet(doc, "Sau khi nhóm xong, giữ lại các biến không đổi trong toàn nhóm.")

    add_body(doc, "9.5. Hàm của đề thi và mạch logic:")
    add_bullet(doc, "Nếu đề mô tả điều kiện đạt/không đạt, hãy đổi trực tiếp điều kiện logic thành hàm Boole.")
    add_bullet(doc, "Sau khi tối thiểu hóa, vẽ mạch bằng cổng NOT, AND, OR theo đúng biểu thức rút gọn.")

    add_table(
        doc,
        ["Dạng biểu diễn", "Cách tạo", "Khi dùng"],
        [
            ["Bảng chân trị -> tổng chuẩn tắc", "Lấy các dòng f = 1, tạo minterm rồi cộng.", "Khi đề cho bảng chân trị."],
            ["Biểu thức -> rút gọn đại số", "Áp dụng luật Boole.", "Khi biểu thức ngắn, dễ biến đổi."],
            ["Bảng Karnaugh", "Nhóm các ô 1 lớn nhất có thể.", "Khi cần biểu thức tối thiểu và mạch logic."],
        ],
        col_widths=[1.5, 2.8, 1.9],
    )

    add_heading(doc, "10. Mẫu xử lý nhanh các dạng bài trong main.pdf", 1)
    add_number(doc, "Cho ma trận kề: xác định vô hướng/hướng, trọng số/không trọng số, liệt kê cạnh, tính bậc, kiểm tra liên thông.")
    add_number(doc, "DFS/BFS: ghi danh sách kề theo thứ tự chữ cái, bắt đầu từ đỉnh yêu cầu, trình bày cây duyệt và thứ tự thăm.")
    add_number(doc, "Euler: đếm bậc hoặc bán bậc, so với điều kiện, rồi kết luận có/không đường đi Euler hoặc chu trình Euler.")
    add_number(doc, "MST: với Kruskal, sắp xếp cạnh; với Prim, chọn cạnh nhỏ nhất nối ra ngoài cây.")
    add_number(doc, "Đường đi ngắn nhất: nếu không trọng số âm thì ưu tiên Dijkstra; nếu có cạnh âm thì dùng Bellman-Ford; nếu mọi cặp đỉnh thì dùng Floyd.")
    add_number(doc, "Luồng cực đại: lập đồ thị dư, tìm đường tăng luồng, đẩy theo nút cổ chai, lặp đến khi không còn đường tăng.")
    add_number(doc, "Boole: từ bảng chân trị tạo minterm, sau đó tối thiểu hóa bằng luật đại số hoặc Karnaugh.")

    add_heading(doc, "11. Checklist trước khi vào phòng thi", 1)
    for item in [
        "Nhìn ma trận là biết đồ thị hướng hay vô hướng.",
        "Biết ngay bài nào là BFS/DFS, bài nào là Euler, bài nào là MST, bài nào là shortest path, bài nào là flow, bài nào là Boolean.",
        "Luôn ghi rõ công thức đang dùng và điều kiện áp dụng.",
        "Nếu có nhiều đáp án tương đương, chọn theo quy ước chữ cái để kết quả ổn định.",
        "Trình bày gọn: công thức -> bước làm -> kết luận.",
    ]:
        add_bullet(doc, item)

    return doc


def main():
    doc = build_document()
    doc.save(OUT_FILE)
    print(f"Saved: {OUT_FILE.resolve()}")


if __name__ == "__main__":
    main()
